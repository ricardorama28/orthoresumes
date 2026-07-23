#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
validar.py — gate de calidad de UPOME.

Uso:
    python engine/validar.py <ID> [--nivel ultra|ultra_plus] [--estructura] [--json]

Devuelve exit code 0 sólo si el tema PASA LOS DOS VEREDICTOS: numérico y
estructural. Cualquier defecto → exit code != 0, para hacer fallar el build.
Con --estructura se corre y reporta sólo el veredicto estructural.

Diseño:
  * Veredicto NUMÉRICO — métricas de contenido (párrafos, caracteres, tablas,
    pasos, callouts, MCQ, cortas, cuantitativos, referencias, densidad de
    negrita) contadas sobre el Markdown fuente, replicando engine/md.js;
    determinístico y coincidente con engine/contar.js. Incluye los invariantes
    de FORMATO del .docx (A4, Calibri, encabezado, pie y borde de página con el
    color del módulo, verificados con python-docx; el borde es el bug nº 1).
  * Veredicto ESTRUCTURAL — completitud de subaspectos por sección según
    engine/esquema.json (derivado de docs/ESQUEMA-12.md). Reporta, por sección,
    los subaspectos ausentes que no fueron cubiertos ni marcados «No aplica».
    Pesa igual que el numérico: un tema con las métricas en verde pero sin
    «criterios de reducción» NO CUMPLE.
"""
import sys
import os
import re
import json
import glob

# Salida UTF-8 robusta en cualquier shell (Git Bash/cmd usan cp1252 y rompen con emojis)
try:
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
SPEC = json.load(open(os.path.join(HERE, "spec.json"), encoding="utf-8"))
ESQUEMA = json.load(open(os.path.join(HERE, "esquema.json"), encoding="utf-8"))

QUANT_RE = re.compile(
    r"\d+(?:[.,]\d+)?\s?(?:mm|cm|°|%|grados?|semanas?|sem\b|meses|mes\b|años?|días?|dias?|mg|Gy|ml|kg|min|N·?m?|puntos?)",
    re.IGNORECASE,
)
CALLOUT_KINDS = ("trick", "warn", "clave")
DIR_OPEN = re.compile(r'^:::(\w+)(?:\s+"([^"]*)")?\s*$')


# ─────────────────────────── localización de archivos ───────────────────────
def modulo_de(tema_id):
    mod = tema_id.split("-")[0].upper()
    if mod not in SPEC["modulos"]:
        raise SystemExit(f"Módulo desconocido en el ID {tema_id!r}")
    return mod


def carpeta_tema(tema_id):
    mod = modulo_de(tema_id)
    base = os.path.join(ROOT, "contenido", mod)
    if not os.path.isdir(base):
        return None
    for d in sorted(os.listdir(base)):
        if d == tema_id or d.startswith(tema_id + "-"):
            p = os.path.join(base, d)
            if os.path.isdir(p):
                return p
    return None


def secciones_md(dir_tema):
    return sorted(glob.glob(os.path.join(dir_tema, "[0-9][0-9]-*.md")))


# ─────────────────────────── conteo sobre el markdown ───────────────────────
def plano(s):
    return re.sub(r"\*\*(.+?)\*\*", r"\1", s)


def contar_markdown(dir_tema):
    m = dict(parrafos=0, caracteres=0, tablas=0, pasos=0, callouts=0, mcq=0,
             cortas=0, cuantitativos=0, referencias=0, negrita_num=0, negrita_den=0)
    texto_quant = []
    callouts_por_seccion = []  # bool por archivo NN

    for archivo in secciones_md(dir_tema):
        lines = open(archivo, encoding="utf-8").read().replace("\r\n", "\n").split("\n")
        i = 0
        callouts_aqui = 0
        while i < len(lines):
            line = lines[i].rstrip()
            if line.strip() == "":
                i += 1
                continue
            mo = DIR_OPEN.match(line)
            if mo:
                kind = mo.group(1)
                body = []
                i += 1
                while i < len(lines) and lines[i].rstrip() != ":::":
                    body.append(lines[i])
                    i += 1
                i += 1  # cerrar :::
                _contar_directiva(kind, body, m, texto_quant)
                if kind in CALLOUT_KINDS:
                    callouts_aqui += 1
                continue

            # títulos / cita / listas / párrafo
            if line.startswith("### ") or line.startswith("## ") or line.startswith("# "):
                txt = line.lstrip("#").strip()
                m["parrafos"] += 1
                m["caracteres"] += len(plano(txt))
                texto_quant.append(txt)
            elif line.startswith("> "):
                txt = line[2:].strip()
                m["parrafos"] += 1
                m["caracteres"] += len(plano(txt))
                texto_quant.append(txt)
                if re.search(r"\b(19|20)\d{2}\b", txt):
                    m["referencias"] += 1
            elif re.match(r"^[-*] ", line) or re.match(r"^\d+\.\s+", line):
                txt = re.sub(r"^([-*]|\d+\.)\s+", "", line).strip()
                m["parrafos"] += 1
                m["caracteres"] += len(plano(txt))
                texto_quant.append(txt)
                m["negrita_den"] += 1
                if re.search(r"\*\*(.+?)\*\*", txt):
                    m["negrita_num"] += 1
            else:
                txt = line.strip()
                m["parrafos"] += 1
                m["caracteres"] += len(plano(txt))
                texto_quant.append(txt)
                m["negrita_den"] += 1
                if re.search(r"\*\*(.+?)\*\*", txt):
                    m["negrita_num"] += 1
            i += 1
        callouts_por_seccion.append(callouts_aqui > 0)

    joined = " ".join(texto_quant)
    m["cuantitativos"] = len(QUANT_RE.findall(joined))
    m["callouts_por_seccion_ok"] = all(callouts_por_seccion) if callouts_por_seccion else False
    m["secciones"] = len(callouts_por_seccion)
    m["densidad"] = (m["negrita_num"] / m["negrita_den"]) if m["negrita_den"] else 0.0
    return m


def _contar_directiva(kind, body, m, texto_quant):
    ne = [l.strip() for l in body if l.strip() != ""]
    if kind in CALLOUT_KINDS:
        m["callouts"] += 1
        for l in ne:
            m["parrafos"] += 1
            m["caracteres"] += len(plano(l))
            texto_quant.append(l)
    elif kind == "pasos":
        steps = [re.sub(r"^\d+\.\s+", "", l).strip() for l in ne]
        m["pasos"] += len(steps)
        m["parrafos"] += 1 + len(steps)  # rótulo + pasos
        for l in steps:
            m["caracteres"] += len(plano(l))
            texto_quant.append(l)
    elif kind == "tabla":
        m["tablas"] += 1  # 0 párrafos y 0 caracteres, pero cuenta como tabla
    elif kind == "mcq":
        q = ""
        opciones = 0
        ans = ""
        for l in ne:
            if l.startswith("P:"):
                q = l[2:].strip()
            elif l.startswith("R:"):
                ans = l[2:].strip()
            elif re.match(r"^[a-hA-H]\)", l):
                opciones += 1
                texto_quant.append(re.sub(r"^[a-hA-H]\)\s*", "", l))
        m["mcq"] += 1
        m["parrafos"] += 1 + opciones + (1 if ans else 0)
        m["caracteres"] += len(plano(q)) + len(plano(ans))
        texto_quant.append(q)
        texto_quant.append(ans)
    elif kind == "corta":
        prompt = ""
        ans = ""
        for l in ne:
            if l.startswith("R:"):
                ans = l[2:].strip()
            elif not prompt:
                prompt = l
        m["cortas"] += 1
        m["parrafos"] += 1 + (1 if ans else 0)
        m["caracteres"] += len(plano(prompt)) + len(plano(ans))
        texto_quant.append(prompt)
        texto_quant.append(ans)
    else:
        # desconocida → como callout
        m["callouts"] += 1
        for l in ne:
            m["parrafos"] += 1
            m["caracteres"] += len(plano(l))
            texto_quant.append(l)


def contar_anki(dir_tema):
    p = os.path.join(dir_tema, "anki.tsv")
    if not os.path.exists(p):
        return 0
    n = 0
    for l in open(p, encoding="utf-8"):
        if l.strip() and not l.startswith("#"):
            n += 1
    return n


# ─────────────────────────── verificación de formato (docx) ─────────────────
def qn(tag):
    from docx.oxml.ns import qn as _qn
    return _qn(tag)


def verificar_formato(docx_path, color_mod):
    """Devuelve (ok:bool, problemas:list, datos:dict) sobre el .docx compilado."""
    problemas = []
    datos = {}
    try:
        import docx
    except ImportError:
        return False, ["python-docx no está instalado (pip install python-docx)"], datos

    if not os.path.exists(docx_path):
        return False, [f"no existe el .docx compilado ({os.path.relpath(docx_path, ROOT)}); corré build.js primero"], datos

    d = docx.Document(docx_path)
    sec = d.sections[0]

    # A4 (tolerancia ±2 mm en EMU; 1 mm = 36000 EMU)
    w, h = int(sec.page_width), int(sec.page_height)
    datos["ancho_mm"] = round(w / 36000, 1)
    datos["alto_mm"] = round(h / 36000, 1)
    if abs(w - 210 * 36000) > 2 * 36000 or abs(h - 297 * 36000) > 2 * 36000:
        problemas.append(f"tamaño de página no es A4 ({datos['ancho_mm']}×{datos['alto_mm']} mm)")

    # Encabezado y pie
    hd = " ".join(p.text for p in sec.header.paragraphs)
    ft = " ".join(p.text for p in sec.footer.paragraphs)
    if "TRAUMATOLOGÍA Y ORTOPEDIA" not in hd.upper():
        problemas.append("el encabezado no dice «TRAUMATOLOGÍA Y ORTOPEDIA»")
    if "VZ" not in ft.upper() or "2026" not in ft:
        problemas.append("el pie no dice «VZ – 2026»")

    # Fuente por defecto Calibri (docDefaults/rPrDefault/rFonts)
    fuente = None
    try:
        styles_el = d.styles.element
        rfonts = styles_el.find(qn("w:docDefaults") + "/" + qn("w:rPrDefault") + "/" + qn("w:rPr") + "/" + qn("w:rFonts"))
        if rfonts is not None:
            fuente = rfonts.get(qn("w:ascii"))
    except Exception:
        pass
    datos["fuente"] = fuente
    if fuente and fuente != "Calibri":
        problemas.append(f"la fuente por defecto no es Calibri (es {fuente})")

    # Borde de página (BUG HISTÓRICO 1): sectPr/pgBorders con color del módulo
    sectPr = sec._sectPr
    pgb = sectPr.find(qn("w:pgBorders"))
    if pgb is None:
        problemas.append("FALTA el borde de página (pgBorders) — bug histórico nº 1; verificá pageBorders{display,offsetFrom}")
    else:
        colores = set()
        for lado in ("top", "left", "bottom", "right"):
            b = pgb.find(qn("w:" + lado))
            if b is not None and b.get(qn("w:color")):
                colores.add(b.get(qn("w:color")).upper())
        datos["borde_colores"] = sorted(colores)
        if color_mod.upper() not in colores:
            problemas.append(f"el borde de página no usa el color del módulo #{color_mod} (encontrado: {sorted(colores) or 'ninguno'})")

    # Conteos del docx (cross-check informativo)
    datos["parrafos_docx"] = sum(1 for p in d.paragraphs if p.text.strip())
    datos["tablas_docx"] = len(d.tables)

    return (len(problemas) == 0), problemas, datos


# ─────────────────────────── veredicto estructural ──────────────────────────
# Contexto mínimo (caracteres además del propio término) para que un término
# cuente como CONTENIDO real y no como un encabezado o un ítem-stub vacío.
MIN_CONTEXTO = 15


def _lineas_seccion(archivo):
    """
    Devuelve (cuerpo, crudo) en minúsculas.
      * cuerpo: líneas de CONTENIDO — se excluyen los títulos (#, ##, ###) para
        que un encabezado vacío no cuente. Incluye prosa, listas, filas de
        tabla, líneas de callout, pasos y marcadores de directiva.
      * crudo: todo el texto en minúsculas, usado sólo para detectar marcadores
        de directiva (:::mcq, :::corta, :::pasos), que denotan un bloque real
        aunque su línea marcadora sea corta.
    """
    crudo = open(archivo, encoding="utf-8").read().replace("\r\n", "\n")
    cuerpo = []
    for ln in crudo.split("\n"):
        s = ln.strip()
        if not s or s.startswith("#"):
            continue
        cuerpo.append(s.lower())
    return cuerpo, crudo.lower()


def _cubierto(sub, cuerpo, crudo):
    """
    Un subaspecto está cubierto si:
      - alguno de sus 'any' es un marcador de directiva (':::…') presente en el
        texto crudo, o
      - alguno de sus 'any' aparece en una línea de CONTENIDO con contexto
        suficiente (la línea tiene ≥ MIN_CONTEXTO caracteres además del término).
    Así, un título vacío o un ítem-stub («- Escalón») NO alcanza para cubrir.
    """
    for kw in sub["any"]:
        k = kw.lower()
        if k.startswith(":::"):
            if k in crudo:
                return True
            continue
        for ln in cuerpo:
            if k in ln and (len(ln) - len(k)) >= MIN_CONTEXTO:
                return True
    return False


def verificar_estructura(dir_tema):
    """
    Recorre los subaspectos de engine/esquema.json por sección y reporta los
    ausentes. La cobertura se evalúa sobre el CONTENIDO real (no los títulos):
    ver _cubierto. La marca «No aplica en esta entidad», al escribirse como
    contenido que incluye el término del subaspecto, cuenta como cubierta.

    @returns dict con: ok (bool), secciones (lista por NN), total, cubiertos,
             ausentes_total.
    """
    esquema = ESQUEMA["secciones"]
    # mapear NN -> ruta de archivo de esa sección
    archivos = {}
    for f in secciones_md(dir_tema):
        nn = os.path.basename(f)[:2]
        archivos[nn] = f

    secciones = []
    total = 0
    cubiertos = 0
    ausentes_total = 0
    for nn in sorted(esquema.keys()):
        subs = esquema[nn]
        total += len(subs)
        if nn not in archivos:
            secciones.append(dict(nn=nn, presentes=0, total=len(subs),
                                  ausentes=[s["label"] for s in subs], falta_archivo=True))
            ausentes_total += len(subs)
            continue
        cuerpo, crudo = _lineas_seccion(archivos[nn])
        ausentes = []
        for s in subs:
            if _cubierto(s, cuerpo, crudo):
                cubiertos += 1
            else:
                ausentes.append(s["label"])
        ausentes_total += len(ausentes)
        secciones.append(dict(nn=nn, presentes=len(subs) - len(ausentes),
                              total=len(subs), ausentes=ausentes, falta_archivo=False))

    ok = ausentes_total == 0
    return dict(ok=ok, secciones=secciones, total=total,
                cubiertos=cubiertos, ausentes_total=ausentes_total)


# ─────────────────────────── evaluación de umbrales ─────────────────────────
def evaluar(m, anki, formato_ok, nivel):
    th = SPEC["niveles"][nivel]
    neg = SPEC["negrita"]
    checks = []

    def chk(nombre, real, req, ok):
        checks.append(dict(metrica=nombre, real=real, requerido=req, ok=ok))

    chk("Párrafos de cuerpo", m["parrafos"], f">= {th['parrafos']}", m["parrafos"] >= th["parrafos"])
    chk("Caracteres", m["caracteres"], f">= {th['caracteres']}", m["caracteres"] >= th["caracteres"])
    chk("Tablas", m["tablas"], f">= {th['tablas']}", m["tablas"] >= th["tablas"])
    chk("Pasos quirúrgicos", m["pasos"], f">= {th['pasos']}", m["pasos"] >= th["pasos"])
    chk("Callouts", m["callouts"], f">= {th['callouts']}", m["callouts"] >= th["callouts"])
    if th["callouts_por_seccion"]:
        chk("Callouts ≥1 por sección", "sí" if m["callouts_por_seccion_ok"] else "no", "todas", m["callouts_por_seccion_ok"])
    chk("MCQ con clave", m["mcq"], f">= {th['mcq']}", m["mcq"] >= th["mcq"])
    chk("Preguntas cortas", m["cortas"], f">= {th['cortas']}", m["cortas"] >= th["cortas"])
    chk("Tarjetas Anki", anki, f">= {th['anki']}", anki >= th["anki"])
    chk("Criterios cuantitativos", m["cuantitativos"], f">= {th['cuantitativos']}", m["cuantitativos"] >= th["cuantitativos"])
    chk("Referencias con año", m["referencias"], f">= {th['referencias']}", m["referencias"] >= th["referencias"])
    dens = m["densidad"]
    chk("Densidad de negrita", f"{dens*100:.0f}%", f"{int(neg['min']*100)}–{int(neg['max']*100)}%", neg["min"] <= dens <= neg["max"])
    chk("Formato .docx", "ok" if formato_ok else "defectos", "invariante", formato_ok)

    todos_ok = all(c["ok"] for c in checks)
    return checks, todos_ok


def nivel_alcanzado(m, anki, formato_ok):
    for niv in ("ultra_plus", "ultra"):
        _, ok = evaluar(m, anki, formato_ok, niv)
        if ok:
            return niv
    return None


def imprimir_estructura(estr):
    """Imprime el bloque del veredicto estructural."""
    print("\n🧩 Completitud estructural (subaspectos de docs/ESQUEMA-12.md)\n")
    nombres = {
        "00": "Definición", "01": "Presentación", "02": "Imágenes", "03": "Diagnóstico",
        "04": "Clasificación", "05": "Epidemiología", "06": "Etiología", "07": "Anatomía",
        "08": "Fisiopatología", "09": "Historia natural", "10": "Tratamiento",
        "11": "Rehabilitación", "12": "Complicaciones", "13": "Cierre",
    }
    for s in estr["secciones"]:
        completo = not s["ausentes"] and not s["falta_archivo"]
        marca = "✅" if completo else "❌"
        etiqueta = f"{s['nn']} {nombres.get(s['nn'], '')}"
        print(f"   {marca}  {etiqueta:<24} {s['presentes']}/{s['total']} subaspectos")
        if s["falta_archivo"]:
            print(f"          • (falta el archivo de la sección {s['nn']})")
        for a in s["ausentes"]:
            print(f"          • ausente: {a}")


# ─────────────────────────── main ───────────────────────────────────────────
def main():
    args = sys.argv[1:]
    tema_id = next((a for a in args if not a.startswith("--")), None)
    if not tema_id:
        print("uso: python engine/validar.py <ID> [--nivel ultra|ultra_plus] [--estructura] [--json]")
        sys.exit(2)
    nivel = "ultra_plus"
    if "--nivel" in args:
        nivel = args[args.index("--nivel") + 1]
    as_json = "--json" in args
    solo_estructura = "--estructura" in args

    dir_tema = carpeta_tema(tema_id)
    if not dir_tema:
        print(f"ERROR: no existe la carpeta de contenido de {tema_id}")
        sys.exit(2)

    # ── Veredicto ESTRUCTURAL (siempre se computa) ──
    estr = verificar_estructura(dir_tema)

    # Modo sólo-estructura: para el redactor, sin depender del .docx compilado.
    if solo_estructura:
        if as_json:
            print(json.dumps(dict(id=tema_id, estructura=estr), ensure_ascii=False, indent=2))
        else:
            print(f"\n🔬 Validación estructural — {tema_id}")
            imprimir_estructura(estr)
            print()
            if estr["ok"]:
                print(f"   ✔ ESTRUCTURA COMPLETA — {estr['cubiertos']}/{estr['total']} subaspectos.\n")
            else:
                print(f"   ✘ ESTRUCTURA INCOMPLETA — faltan {estr['ausentes_total']} subaspectos "
                      f"(marcá los que no aplican con «No aplica en esta entidad»).\n")
        sys.exit(0 if estr["ok"] else 1)

    # ── Veredicto NUMÉRICO ──
    color = SPEC["modulos"][modulo_de(tema_id)]["color"]
    m = contar_markdown(dir_tema)
    anki = contar_anki(dir_tema)
    docx_path = os.path.join(ROOT, "salida", "docx", f"{tema_id}.docx")
    formato_ok, problemas, datos = verificar_formato(docx_path, color)

    checks, num_ok = evaluar(m, anki, formato_ok, nivel)
    alcanzado = nivel_alcanzado(m, anki, formato_ok)
    ok = num_ok and estr["ok"]

    if as_json:
        print(json.dumps(dict(id=tema_id, nivel_objetivo=nivel, nivel_alcanzado=alcanzado,
                              ok=ok, numerico_ok=num_ok, estructura_ok=estr["ok"],
                              checks=checks, estructura=estr,
                              formato=dict(ok=formato_ok, problemas=problemas, datos=datos)),
                         ensure_ascii=False, indent=2))
        sys.exit(0 if ok else 1)

    print(f"\n🔬 Validación — {tema_id}   (objetivo: {nivel})")
    print("\n📏 Veredicto numérico\n")
    for c in checks:
        marca = "✅" if c["ok"] else "❌"
        print(f"   {marca}  {c['metrica']:<26} {str(c['real']):>10}   (req {c['requerido']})")
    if problemas:
        print("\n   Formato — defectos:")
        for p in problemas:
            print(f"      • {p}")
    if datos:
        extra = []
        if "borde_colores" in datos:
            extra.append("borde=" + ",".join(datos["borde_colores"]))
        if "fuente" in datos and datos["fuente"]:
            extra.append("fuente=" + datos["fuente"])
        if "parrafos_docx" in datos:
            extra.append(f"párrafos(docx)={datos['parrafos_docx']}")
        if extra:
            print("\n   docx: " + " · ".join(extra))

    # ── Veredicto ESTRUCTURAL ──
    imprimir_estructura(estr)

    # ── Resumen ──
    print("\n" + "   " + "─" * 56)
    alc = alcanzado.upper() if alcanzado else "ninguno (por debajo de ULTRA)"
    marca_num = "✔" if num_ok else "✘"
    marca_est = "✔" if estr["ok"] else "✘"
    print(f"   {marca_num} Numérico: {'OK — nivel ' + alc if num_ok else 'NO alcanza ' + nivel.upper() + ' (actual: ' + alc + ')'}")
    print(f"   {marca_est} Estructural: {'COMPLETA' if estr['ok'] else 'INCOMPLETA — faltan ' + str(estr['ausentes_total']) + ' subaspectos'}")
    print()
    if ok:
        print(f"   ✔ TODAS LAS VALIDACIONES OK — nivel alcanzado: {alcanzado.upper()}\n")
        sys.exit(0)
    else:
        print(f"   ✘ NO CUMPLE {nivel.upper()} (requiere numérico Y estructural en verde)\n")
        sys.exit(1)


if __name__ == "__main__":
    main()
